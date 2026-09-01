from docx import Document
from docx.shared import Pt
import os

doc = Document()

# Title
title = doc.add_heading('LegalMetro AI \nAutomated Compliance Verification System', 0)

# Section 1
doc.add_heading('1. Title', level=1)
doc.add_paragraph('LegalMetro AI — Automated Compliance Verification System for Packaged Commodities under the Legal Metrology (Packaged Commodities) Rules, 2011')

# Section 2
doc.add_heading('2. Understanding the Problem', level=1)
doc.add_paragraph(
    "Every packaged commodity sold in India must carry a fixed set of mandatory declarations under the Legal Metrology Act, 2009 and the Legal Metrology (Packaged Commodities) Rules, 2011. These include:\n"
    "- Name and address of the manufacturer/packer/importer\n"
    "- Net quantity (in standard units)\n"
    "- Maximum Retail Price (MRP) inclusive of all taxes\n"
    "- Month and year of manufacture/packing/import\n"
    "- Consumer care/complaint details\n"
    "- Common/generic name of the commodity\n"
    "- Unit sale price (where applicable)\n"
    "- Country of origin (for imported goods)\n\n"
    "Enforcement agencies currently rely on manual, visual inspection of these declarations — a process that is slow, inconsistent, and cannot scale to the sheer volume of SKUs across physical retail and e-commerce. Common violations (missing declarations, undersized fonts, incorrect MRP formatting, absent manufacturer details) go undetected simply because there aren't enough inspectors to check every product.\n"
    "The problem, therefore, is not a lack of rules but a lack of scalable, technology-driven enforcement tooling."
)

# Section 3
doc.add_heading('3. Proposed Solution — Overview', level=1)
doc.add_paragraph(
    "LegalMetro AI is a web + mobile platform that uses computer vision (OCR) and rule-based validation to automatically:\n"
    "1. Scan an image of a product label/package.\n"
    "2. Extract every text/graphic element on it.\n"
    "3. Map extracted content against the mandatory declaration checklist prescribed in the 2011 Rules.\n"
    "4. Validate correctness, placement, and font-size/readability compliance.\n"
    "5. Flag violations and auto-generate a compliance report.\n"
    "6. Store the record in a searchable repository for enforcement officers and provide analytics dashboards.\n\n"
    "In short: 'Point camera -> scan label -> get compliance report in seconds.'"
)

# Section 4
doc.add_heading('4. Technical Stack Implemented', level=1)
doc.add_paragraph(
    "- Frontend: React.js (Vite) + Tailwind CSS (Beautiful Enforcement Dashboard)\n"
    "- Backend: Python FastAPI (REST API, robust and scalable)\n"
    "- Database: PostgreSQL (Relational schema for Inspections and Products)\n"
    "- AI / OCR Layer: OpenCV (Image preprocessing) + Tesseract OCR (Text Extraction)\n"
    "- Rule Engine: Custom Python Logic mapping extracted JSON to Legal Metrology 2011 Rules\n"
    "- Report Generation: jsPDF (Client-side automated PDF generation)"
)

# Section 5
doc.add_heading('5. System Workflow', level=1)
doc.add_paragraph(
    "1. Inspector logs in (web dashboard) -> selects 'New Scan'.\n"
    "2. Captures/uploads images of the product.\n"
    "3. Backend pre-processes image -> runs OCR -> extracts text.\n"
    "4. NLP/Regex extracts MRP, Date, and Quantity.\n"
    "5. Rule engine cross-checks extracted data against LM Rules, 2011 ruleset.\n"
    "6. Dashboard instantly displays a compliance report highlighting compliant vs. violated fields (Red/Green badges).\n"
    "7. Inspector downloads the official PDF Challan/Report for record."
)

doc.save('LegalMetro_AI_Project_Proposal.docx')
print("Document saved successfully!")
